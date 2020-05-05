#!/usr/bin/env python3

# PYTHON_ARGCOMPLETE_OK

import os,sys,re,json,codecs,xmltodict,unicodedata

from collections import OrderedDict
from io import StringIO

from Perdy.pyxbext import directory

from xlrd import open_workbook,xldate_as_tuple
from xlwt import Workbook

from docx import Document
from docx.styles.style import _ParagraphStyle

from Argumental.Argue import Argue
from Baubles.Logger import Logger
from Perdy.parser import printXML
from Perdy.pretty import prettyPrintLn, Style
from GoldenChild.xpath import *

args=Argue()
logger=Logger()

#_________________________________________________________________
@args.command(single=True)
class OPML(object):
	'''
	opml conversion tool
	'''

	#_____________________________________________________________
	def _clean(self, text):
		if not text:
			return ''
		if type(text) is str:
			text = unicode(text)
		return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore')

	#.............................................................
	def _load(self, file, extension=None):
		(self.doc,self.ctx,self.nsp) = getContextFromFile(file)
		if extension:
			backup = '%s.%s'%(file,extension)
			if os.path.isfile(backup):
				os.unlink(backup)
			os.rename(file, backup)

	#.............................................................
	def _save(self, file):
		with codecs.open(file, 'w', encoding='utf8') as output:
			printXML(str(self.doc), output = output)
			
	#.............................................................
	@args.operation
	def opml2org(self, file):
		'''
		from cloud outliner to emacs org mode
		'''
		self._load(file,'outline')
		xH = '//outline[@text]'
		for headline in getElements(self.ctx,xH):
			setAttribute(headline, 'structure', 'headline')
			note = getAttribute(headline, '_note')
			delAttribute(headline, '_note')
			if not note: continue
			for paragraph in note.split('&#10;'):
				child = addElement(self.doc,'outline',headline)
				setAttribute(child,'structure','paragraph')
				setAttribute(child,'text', paragraph)
		self._save(file)

	#.............................................................
	def _recurse(self, sheet, parent, row=0, col=0):
		for child in getElements(self.ctx, 'outline', parent):
			try:
				text = getAttribute(child, 'text')
				text = self._clean(text)
			except:
				text = '?'
			#print(row, col, text)
			sheet.write(row,col,text)
			try:
				note = getAttribute(child, '_note')
				node = self._clean(node)
			except:
				note = None
			if note:
				row += 1
				sheet.write(row,col+1,note)
			row = self._recurse(sheet, child, row+1, col+1)
		return row

	#.............................................................
	@args.operation
	def opml2excel(self, file):
		'''
		convert opml to excel format
		'''
		if not file.endswith('opml'):
			sys.stderr.write('not an opml file = %s\n'%file)
			return

		self._load(file)
		workbook = Workbook()
		sheet = workbook.add_sheet('opml')
		self._recurse(sheet, getElement(self.ctx, '/opml/body'))
		workbook.save(file.replace('.opml','.xls'))

	#.............................................................
	@args.operation
	def org2opml(self, file):
		'''
		from emacs org mode to cloud outliner
		'''
		self._load(file,'org')
		xH = '//outline[@structure="headline"]'
		xP = 'outline[@structure="paragraph"]'
		for headline in getElements(self.ctx,xH):
			delAttribute(headline, 'structure')
			sio = StringIO()
			for paragraph in getElements(self.ctx,xP,headline):
				sio.write('%s\n'%paragraph.prop('text'))
				paragraph.unlinkNode()
				paragraph.freeNode()
			note = sio.getvalue()
			setAttribute(headline, '_note', note)
		self._save(file)

	#.............................................................
	@args.operation
	def text2opml(self, file):
		'''
		read a text file using the indent as the level
		'''
		textPattern   = re.compile('^(\s+)(\S\S.*)$')
		bulletPattern = re.compile('^(\S[\.\)])(\s*)(\S.*)$')
		
		with open(file) as input:
			for line in input.read().split('\n'):
				line = line.rstrip('\n')
				textMatch = textPattern.match(line)
				if textMatch:
					(indent, text) = textMatch.groups()
				else:
					bulletMatch = bulletPattern.match(line)
					if bulletMatch:
						(indent, text) = bulletMatch.groups()[1:]
					else:
						continue
				print('%s%s\n'%(indent, text))

	#.............................................................
	@args.operation
	def json2opml(self, file):
		'''
		read a json file and convert to opml
		'''

		ext = '.%s'%(file.split('.')[-1].lower())
		if ext != '.json':
			sys.stderr.write('not a json file %s\n'%file)
			return
		
		def dig(node, parent):
			if type(node) in [list]:
				parent['outline'] = []
				for i in range(len(node)):
					item = node[i]
					child = OrderedDict([
						('@text','[%s]'%i)
					])
					dig(item, child)
					parent['outline'].append(child)
				return 
			
			if type(node) in [dict,OrderedDict]:
				parent['outline'] = []
				for name, value in node.items():
					child = OrderedDict([
						('@text', name),
					])
					dig(value, child)
					parent['outline'].append(child)
				return
			
			# assume fundamental fromhere
			parent['@_note'] = node


		opml = OrderedDict([
			('opml', OrderedDict([
				('@version', '1.0'),
				('head', OrderedDict([
					('title', file),
				])),
				('body', OrderedDict([
				])) 
			]))
		])
		
		with open(file) as input:
			d = json.load(input)
			dig(d, opml['opml']['body'])
			#print(json.dumps(opml,indent=4))
				  
			with open(file.replace(ext,'.opml'),'w') as output:
				xmltodict.unparse(opml, output)
					  
	#.............................................................
	@args.operation
	def docxHeadings2opml(self, file):
		if not file.endswith('docx'):
			sys.stderr.write('not a docx file\n')
			return
		doc = Document(file)

		opml = OrderedDict([
			('opml', OrderedDict([
				('@version', '1.0'),
				('head', OrderedDict([
					('title', file),
				])),
				('body', OrderedDict([
				])) 
			]))
		])
		
		body = opml['opml']['body']
		# stack[-1] is the current parent
		stack = [ body ]

		for paragraph in doc.paragraphs:
			ps = paragraph.style
			text = paragraph.text.strip().replace('\t',' ').replace('\n',' ')

			if len(text) == 0: continue
			
			if ps.name.startswith('Normal'):
				print('%s   "%s"'%('  '*(level-1), text))
				
				parent = stack[-1]['outline'][-1]
				if '@_note' not in parent.keys():
					parent['@_note'] = ''
				lines = list(filter(lambda x: len(x), parent['@_note'].split('\n')))
				lines.append(text)
				parent['@_note'] = '\n'.join(lines)
				continue
			
			if not ps.name.startswith('Heading'):
				#sys.stderr.write('%s\n'%ps.name)
				continue

			heading = ps.name.replace('Heading ','').split(' ')[0]
			level=int(heading)

			print('%s %s'%('  '*(level-1), text))

			outline = {
				'@text': text,
			}

			if level > len(stack):
				parent = stack[-1]['outline'][-1]
				stack.append(parent)

			if level < len(stack):
				stack.pop()
				
			if 'outline' not in stack[-1].keys():
				stack[-1]['outline'] = []
			stack[-1]['outline'].append(outline)				

			
		name = '.outline.opml'
		with open(name, 'w') as output:
			xmltodict.unparse(opml, output, pretty=True)
		print(name)
		
		name = '.outline.json'
		with open(name, 'w') as output:
			json.dump(opml, output, indent=4)
		print(name)

		
	#.............................................................
	@args.operation
	def docxIndented2opml(self, file):
		if not file.endswith('docx'):
			sys.stderr.write('not a docx file\n')
			return
		doc = Document(file)

		opml = OrderedDict([
			('opml', OrderedDict([
				('@version', '1.0'),
				('head', OrderedDict([
					('title', file),
				])),
				('body', OrderedDict([
				])) 
			]))
		])
		
		body = opml['opml']['body']
		# stack[-1] is the current parent
		stack = [ body ]

		p = re.compile('^(\t*| *)(\S.*)$')
		
		for paragraph in doc.paragraphs:
			ps = paragraph.style
			pf = paragraph.paragraph_format

			text = paragraph.text.strip()
			
			if pf.first_line_indent or pf.left_indent:
				print(paragraph.text)

			if len(text) == 0: continue

			level=1

			outline = {
				'@text': text,
			}

			if level > len(stack):
				parent = stack[-1]['outline'][-1]
				stack.append(parent)

			if level < len(stack):
				stack.pop()
				
			if 'outline' not in stack[-1].keys():
				stack[-1]['outline'] = []
			stack[-1]['outline'].append(outline)				

			
		name = file.replace('docx','opml')
		with open(name,'w') as output:
			xmltodict.unparse(opml, output)
		print(name)

		name = file.replace('docx','json')
		with open(name,'w') as output:
			json.dump(opml, output, indent=4)


#_________________________________________________________________
if __name__ == '__main__': args.execute()

